import base64
import copy
import html
import io
import json
import os
import re

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openai import OpenAI, OpenAIError, RateLimitError
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Flowable, HRFlowable, Paragraph, SimpleDocTemplate, Spacer


# Current app stage:
# 1. OCR + translation: done
# 2. Structured extraction with layout metadata: done in this version
# 3. Cleaner A4 layout recreation: done in this version as an HTML/CSS preview
# 4. PDF export: done in this version
#
# Not implemented yet:
# 5. Canva-style design polish

MODEL = "gpt-5.5"
IMAGE_DETAIL = "original"
ELEMENT_TYPES = [
    "title",
    "subtitle",
    "heading",
    "body_block",
    "callout",
    "label",
    "diagram_text",
    "chain_diagram",
    "footer",
    "list_item",
    "step",
    "subitem",
    "bullet_item",
    "divider",
]
LAYOUT_HINTS = [
    "centered",
    "left",
    "right",
    "indented",
    "small",
    "large",
    "bold",
    "uppercase",
    "boxed",
    "diagram-label",
    "numbered",
    "lettered",
    "bullet",
    "compact",
    "top-rule",
    "bottom-rule",
    "section-break",
    "chain-diagram",
]

DEFAULT_CHAIN_LABELS = {
    "vulnerability": "RANLJIVOST",
    "prompting_event": "SPROŽILNI DOGODEK",
    "links": "ČLENI",
    "consequences": "POSLEDICE",
    "problem_behavior": "PROBLEMATIČNO VEDENJE",
}

CHAIN_LABEL_MATCHES = {
    "vulnerability": ["vulnerability", "ranljivost"],
    "prompting_event": ["prompting event", "sprožilni dogodek"],
    "links": ["links", "členi"],
    "consequences": ["consequences", "posledice"],
    "problem_behavior": ["problem behavior", "problematično vedenje"],
}


def get_api_key():
    """Read the OpenAI API key from Streamlit secrets or Windows environment."""
    try:
        key_from_secrets = st.secrets.get("OPENAI_API_KEY", "")
    except Exception:
        key_from_secrets = ""

    key = key_from_secrets or os.getenv("OPENAI_API_KEY", "")
    if key == "PASTE_YOUR_OPENAI_API_KEY_HERE":
        return ""
    return key


def uploaded_image_to_data_url(uploaded_file):
    """Convert the uploaded image into a base64 data URL for the OpenAI API."""
    image_bytes = uploaded_file.getvalue()
    image_base64 = base64.b64encode(image_bytes).decode("utf-8")
    mime_type = uploaded_file.type or "image/jpeg"
    return f"data:{mime_type};base64,{image_base64}"


def build_prompt():
    """Ask the model for structured extraction, translation, and layout hints."""
    return """
You are helping build a personal app that translates scanned English psychotherapy
handouts into Slovenian and recreates a similar handout layout.

This app recreates the page from your JSON. Return only structured content.
Do not generate PDF, HTML, or Word instructions.

Tasks:
1. Read all visible English text in the image.
2. Extract the handout structure as JSON, including visual layout details.
3. Translate each element into Slovenian.
4. Preserve the order of elements.
5. Add layout data so an A4 HTML/CSS preview can be recreated.

Use these preferred Slovenian psychotherapy/DBT terms:
- chain analysis = verižna analiza
- problem behavior = problematično vedenje
- prompting event = sprožilni dogodek
- vulnerability factors = dejavniki ranljivosti
- links = členi
- consequences = posledice
- skillful behavior = vešče vedenje
- prevention strategy/plans = strategija preprečevanja / načrti preprečevanja
- repair = popraviti / popravite

Return only valid JSON in this exact shape:
{
  "page": {
    "layout_profile": "worksheet | handout | diagram-heavy | dense-text",
    "title": {
      "english_text": "",
      "slovenian_text": ""
    },
    "subtitle": {
      "english_text": "",
      "slovenian_text": ""
    },
    "footer": {
      "english_text": "",
      "slovenian_text": ""
    },
    "elements": [
      {
        "order": 1,
        "element_type": "title | subtitle | heading | body_block | callout | label | diagram_text | chain_diagram | footer | list_item | step | subitem | bullet_item | divider",
        "english_text": "",
        "slovenian_text": "",
        "layout_hint": "space-separated fallback hints such as centered bold bottom-rule numbered compact",
        "layout": {
          "alignment": "left | center | right",
          "text_size": "title | subtitle | heading | body | small | tiny",
          "emphasis": "normal | bold | uppercase | italic",
          "indent_level": 0,
          "list_marker": "",
          "marker_style": "none | number | letter | bullet",
          "spacing": "normal | compact | section-break",
          "border": "none | top-rule | bottom-rule | boxed",
          "group": "header | diagram | body | footer"
        },
        "diagram": {
          "diagram_type": "none | chain_analysis",
          "labels": {
            "vulnerability": "RANLJIVOST",
            "prompting_event": "SPROŽILNI DOGODEK",
            "links": "ČLENI",
            "consequences": "POSLEDICE",
            "problem_behavior": "PROBLEMATIČNO VEDENJE"
          }
        }
      }
    ]
  },
  "structured": {
    "title": [],
    "subtitle": [],
    "headings": [],
    "body_blocks": [],
    "callouts": [],
    "labels": [],
    "diagram_text": [],
    "diagrams": [],
    "footer": []
  },
  "slovenian_sections": [
    {
      "title": "short Slovenian section title",
      "text": "translated Slovenian text"
    }
  ],
  "unclear_text": [
    "any words or lines that were hard to read"
  ]
}

Important:
- Keep all translated content clinically careful and natural in Slovenian.
- Do not summarize away details.
- If the original has numbered steps, use element_type "step".
- If the original has A/B/C lettered subitems, use element_type "subitem".
- If the original has bullet points, use element_type "bullet_item".
- Put exact visible markers such as "1.", "A.", or "•" in layout.list_marker.
- Do not repeat the marker at the start of slovenian_text unless you are unsure.
- Never omit an important visual diagram.
- If the original has the chain-analysis image with ovals/links, create one element with element_type "chain_diagram".
- For a chain-analysis diagram, put translated labels in diagram.labels using the keys vulnerability, prompting_event, links, consequences, and problem_behavior.
- Use layout_hint "centered chain-diagram section-break" and layout.group "diagram" for a chain-analysis diagram.
- Use element_type "diagram_text" only for standalone diagram labels that are not already represented inside a chain_diagram.
- If a line is visually a header, title, subtitle, footer, label, or callout, mark it that way.
- Use layout.border "bottom-rule" for title/header lines that have an underline.
- Use layout.spacing "section-break" before major sections.
""".strip()


def call_openai_vision(image_data_url):
    """Send the image and prompt to OpenAI and return the model text."""
    client = OpenAI(api_key=get_api_key())

    response = client.responses.create(
        model=MODEL,
        input=[
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": build_prompt()},
                    {
                        "type": "input_image",
                        "image_url": image_data_url,
                        "detail": IMAGE_DETAIL,
                    },
                ],
            }
        ],
    )

    return response.output_text


def call_openai_layout_revision(result, instruction, page_count):
    """Ask GPT-5.5 to revise the structured handout after initial creation."""
    client = OpenAI(api_key=get_api_key())
    prompt = f"""
You are revising structured JSON for an A4 handout recreation app.

Goal:
Improve the existing structured handout so the rendered HTML/PDF/DOCX layout is cleaner and less likely to be cut off.

User instruction:
{instruction}

Current PDF page count from the app: {page_count}

Rules:
- Return only valid JSON in the same schema as the input.
- Preserve all meaning and all translated Slovenian text unless the instruction explicitly asks for wording changes.
- Prefer layout changes over deleting content.
- If content is likely cut off or overflowing, make layout more compact:
  - use layout_hint tokens such as compact, small, section-break, centered, left, boxed, bottom-rule
  - reduce excessive section breaks
  - combine overly fragmented body elements when that helps
  - keep diagrams and labels readable
- Keep the page.elements list in correct visual order.
- Do not add markdown or explanation outside the JSON.

Current structured JSON:
{json.dumps(result, ensure_ascii=False, indent=2)}
""".strip()

    response = client.responses.create(
        model=MODEL,
        input=[
            {
                "role": "user",
                "content": [{"type": "input_text", "text": prompt}],
            }
        ],
    )
    return response.output_text


def parse_model_json(model_text):
    """Parse JSON even if the model accidentally wraps it in a code block."""
    cleaned = model_text.strip()

    code_block = re.search(r"```(?:json)?\s*(.*?)```", cleaned, re.DOTALL)
    if code_block:
        cleaned = code_block.group(1).strip()

    return normalize_result(json.loads(cleaned))


def text_matches_chain_label(text, label_key):
    """Check whether a text fragment is one of the known chain diagram labels."""
    text_lower = text.lower()
    return any(match in text_lower for match in CHAIN_LABEL_MATCHES[label_key])


def chain_label_key_for_text(text):
    """Map a translated or English diagram label to a chain diagram label key."""
    for label_key in CHAIN_LABEL_MATCHES:
        if text_matches_chain_label(text, label_key):
            return label_key
    return ""


def diagram_labels_from_element(element):
    """Read chain diagram labels from the model's optional diagram object."""
    labels = DEFAULT_CHAIN_LABELS.copy()
    diagram = element.get("diagram", {})
    if not isinstance(diagram, dict):
        return labels

    raw_labels = diagram.get("labels", {})
    if isinstance(raw_labels, dict):
        for label_key in DEFAULT_CHAIN_LABELS:
            value = raw_labels.get(label_key, "")
            if isinstance(value, dict):
                value = value.get("slovenian_text") or value.get("text") or ""
            if isinstance(value, str) and value.strip():
                labels[label_key] = value.strip()
    elif isinstance(raw_labels, list):
        for item in raw_labels:
            if not isinstance(item, dict):
                continue
            label_key = item.get("key") or chain_label_key_for_text(
                f"{item.get('english_text', '')} {item.get('slovenian_text', '')}"
            )
            value = item.get("slovenian_text") or item.get("text") or ""
            if label_key in labels and value:
                labels[label_key] = str(value).strip()

    return labels


def find_chain_diagram_from_text_labels(elements):
    """Detect older model output that listed chain labels but omitted the graphic."""
    labels = DEFAULT_CHAIN_LABELS.copy()
    matched_indexes = []
    first_order = None

    for index, element in enumerate(elements):
        if element.get("element_type") != "diagram_text":
            continue

        combined_text = " ".join(
            [
                element.get("english_text", ""),
                element.get("slovenian_text", ""),
            ]
        )
        label_key = chain_label_key_for_text(combined_text)
        if not label_key:
            continue

        translated = element.get("slovenian_text", "").strip()
        if translated:
            labels[label_key] = translated.upper()
        matched_indexes.append(index)
        first_order = element.get("order", first_order) if first_order is None else first_order

    if len(set(matched_indexes)) < 3:
        return None, []

    return {
        "order": first_order or 1,
        "element_type": "chain_diagram",
        "english_text": "Chain analysis diagram",
        "slovenian_text": "",
        "layout_hint": "centered chain-diagram section-break",
        "layout": {
            "alignment": "center",
            "text_size": "body",
            "emphasis": "normal",
            "indent_level": 0,
            "list_marker": "",
            "marker_style": "none",
            "spacing": "section-break",
            "border": "none",
            "group": "diagram",
        },
        "diagram": {
            "diagram_type": "chain_analysis",
            "labels": labels,
        },
    }, matched_indexes


def normalize_result(result):
    """Fill in missing visual structure so important diagrams survive rendering."""
    page = result.setdefault("page", {})
    elements = page.setdefault("elements", [])
    if not isinstance(elements, list):
        page["elements"] = []
        return result

    has_chain_diagram = any(
        element.get("element_type") == "chain_diagram" for element in elements if isinstance(element, dict)
    )
    if not has_chain_diagram:
        chain_diagram, matched_indexes = find_chain_diagram_from_text_labels(elements)
        if chain_diagram:
            kept_elements = [
                element for index, element in enumerate(elements) if index not in matched_indexes
            ]
            kept_elements.append(chain_diagram)
            elements = sorted(kept_elements, key=lambda item: item.get("order", 0))
            for index, element in enumerate(elements, start=1):
                element["order"] = index
            page["elements"] = elements
    else:
        for element in elements:
            if element.get("element_type") == "chain_diagram":
                element.setdefault("diagram", {})
                element["diagram"].setdefault("diagram_type", "chain_analysis")
                element["diagram"]["labels"] = diagram_labels_from_element(element)

    return result


def revise_handout_with_gpt(result, instruction, page_count):
    """Revise the current handout JSON with GPT-5.5 and store it as the latest result."""
    if not instruction.strip():
        st.warning("Describe what GPT-5.5 should improve before running the re-edit.")
        return

    with st.status("Re-editing handout with GPT-5.5...", expanded=True) as status:
        try:
            st.write("Sending current layout JSON and your instruction to GPT-5.5.")
            model_text = call_openai_layout_revision(result, instruction, page_count)
            st.write("Parsing revised structured JSON.")
            revised_result = parse_model_json(model_text)
        except json.JSONDecodeError:
            status.update(label="Finished, but JSON parsing failed.", state="error")
            st.warning("GPT returned text that was not valid JSON. Showing raw output.")
            st.text_area("Raw revision output", model_text, height=400)
            return
        except OpenAIError as error:
            status.update(label="OpenAI API error.", state="error")
            st.error(str(error))
            return

        status.update(label="Re-edit complete.", state="complete")

    st.session_state["latest_result"] = revised_result
    st.session_state["result_version"] = st.session_state.get("result_version", 0) + 1
    st.session_state["handout_revision_instruction"] = instruction
    st.rerun()


def safe_int(value, default=0):
    """Convert model-provided layout numbers without crashing on messy JSON."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def split_layout_hint(layout_hint):
    """Turn a layout hint string into known tokens."""
    if not layout_hint:
        return []
    if isinstance(layout_hint, list):
        raw_tokens = layout_hint
    else:
        raw_tokens = re.split(r"[\s,|/]+", str(layout_hint))
    return [token for token in raw_tokens if token in LAYOUT_HINTS]


def get_element_layout(element):
    """Return the richer layout object if the model supplied one."""
    layout = element.get("layout", {})
    if isinstance(layout, dict):
        return layout
    return {}


def layout_tokens_for_element(element):
    """Combine older layout_hint strings with the newer layout object."""
    tokens = set(split_layout_hint(element.get("layout_hint", "")))
    layout = get_element_layout(element)

    alignment = str(layout.get("alignment", "")).strip().lower()
    if alignment == "center":
        tokens.add("centered")
    elif alignment in {"left", "right"}:
        tokens.add(alignment)

    emphasis = str(layout.get("emphasis", "")).strip().lower()
    if emphasis in {"bold", "uppercase"}:
        tokens.add(emphasis)

    text_size = str(layout.get("text_size", "")).strip().lower()
    if text_size in {"small", "tiny"}:
        tokens.add("small")
    elif text_size in {"title", "heading"}:
        tokens.add("large")

    marker_style = str(layout.get("marker_style", "")).strip().lower()
    if marker_style == "number":
        tokens.add("numbered")
    elif marker_style == "letter":
        tokens.add("lettered")
    elif marker_style == "bullet":
        tokens.add("bullet")

    spacing = str(layout.get("spacing", "")).strip().lower()
    if spacing in {"compact", "section-break"}:
        tokens.add(spacing)

    border = str(layout.get("border", "")).strip().lower()
    if border in {"boxed", "top-rule", "bottom-rule"}:
        tokens.add(border)

    if safe_int(layout.get("indent_level"), 0) > 0:
        tokens.add("indented")

    return [token for token in LAYOUT_HINTS if token in tokens]


def css_token(token):
    """Convert a layout token into a safe CSS class suffix."""
    return token.replace("-", "_")


def format_html_text(text):
    """Escape translated text while keeping simple line breaks visible."""
    return html.escape(text.strip()).replace("\n", "<br>")


def marker_for_element(element, text):
    """Find a numbered, lettered, or bullet marker for worksheet rows."""
    layout = get_element_layout(element)
    marker = str(layout.get("list_marker", "")).strip()
    if marker:
        return marker

    match = re.match(r"^\s*((?:\d+|[A-Za-z])[\.)]|[-\u2022])\s+", text)
    if match:
        return match.group(1)

    element_type = element.get("element_type", "body_block")
    tokens = layout_tokens_for_element(element)
    if element_type == "bullet_item" or "bullet" in tokens:
        return "\u2022"
    return ""


def remove_repeated_marker(text, marker):
    """Avoid showing list markers twice when the model includes them in text."""
    if not marker:
        return text.strip()
    return re.sub(r"^\s*" + re.escape(marker) + r"\s*", "", text.strip(), count=1)


def export_text_for_element(element):
    """Return text for PDF/Word exports, including list markers when needed."""
    raw_text = element.get("slovenian_text", "").strip()
    marker = marker_for_element(element, raw_text)
    if marker:
        return f"{marker} {remove_repeated_marker(raw_text, marker)}".strip()
    return raw_text


def render_chain_diagram(element, classes):
    """Render the DBT chain-analysis image as a clean translated SVG."""
    labels = diagram_labels_from_element(element)
    vulnerability = html.escape(labels["vulnerability"])
    prompting_event = html.escape(labels["prompting_event"])
    links = html.escape(labels["links"])
    consequences = html.escape(labels["consequences"])
    problem_behavior = html.escape(labels["problem_behavior"])

    return f"""
<figure class="{classes} chain-diagram" aria-label="Diagram verižne analize">
  <svg viewBox="0 0 720 180" role="img" aria-label="Diagram verižne analize">
    <line class="chain-arrow" x1="126" y1="70" x2="172" y2="78"></line>
    <polygon class="chain-arrow-head" points="172,78 160,72 160,85"></polygon>

    <line class="chain-arrow" x1="620" y1="54" x2="548" y2="68"></line>
    <polygon class="chain-arrow-head" points="548,68 560,60 562,73"></polygon>

    <line class="chain-arrow" x1="352" y1="132" x2="372" y2="96"></line>
    <polygon class="chain-arrow-head" points="372,96 372,110 360,104"></polygon>
    <line class="chain-arrow" x1="386" y1="132" x2="410" y2="96"></line>
    <polygon class="chain-arrow-head" points="410,96 408,110 397,103"></polygon>

    <line class="chain-arrow" x1="508" y1="126" x2="505" y2="92"></line>
    <polygon class="chain-arrow-head" points="505,92 511,105 499,105"></polygon>

    <ellipse class="chain-oval" cx="208" cy="78" rx="58" ry="25" transform="rotate(10 208 78)"></ellipse>
    <ellipse class="chain-oval" cx="250" cy="94" rx="49" ry="22" transform="rotate(8 250 94)"></ellipse>
    <ellipse class="chain-oval" cx="318" cy="85" rx="39" ry="19" transform="rotate(8 318 85)"></ellipse>
    <ellipse class="chain-oval" cx="378" cy="88" rx="52" ry="25" transform="rotate(7 378 88)"></ellipse>
    <ellipse class="chain-oval" cx="438" cy="84" rx="47" ry="22" transform="rotate(-6 438 84)"></ellipse>
    <ellipse class="chain-oval chain-oval-problem" cx="505" cy="75" rx="78" ry="27" transform="rotate(1 505 75)"></ellipse>
    <ellipse class="chain-oval" cx="590" cy="83" rx="36" ry="19" transform="rotate(8 590 83)"></ellipse>
    <ellipse class="chain-oval" cx="630" cy="101" rx="39" ry="19" transform="rotate(8 630 101)"></ellipse>
    <ellipse class="chain-oval" cx="662" cy="121" rx="41" ry="18" transform="rotate(11 662 121)"></ellipse>

    <text class="chain-label" x="72" y="66" text-anchor="middle">{vulnerability}</text>
    <text class="chain-label" x="214" y="132" text-anchor="middle">{prompting_event}</text>
    <text class="chain-label" x="368" y="151" text-anchor="middle">{links}</text>
    <text class="chain-label" x="505" y="132" text-anchor="middle">{consequences}</text>
    <text class="chain-label" x="632" y="48" text-anchor="middle">{problem_behavior}</text>
  </svg>
</figure>
""".strip()


def element_class(element):
    """Choose a CSS class based on the extracted element type and layout hint."""
    element_type = element.get("element_type", "body_block")
    layout = get_element_layout(element)
    tokens = layout_tokens_for_element(element)

    classes = ["handout-element", f"type-{element_type}"]
    classes.extend(f"hint-{css_token(token)}" for token in tokens)

    if "centered" in tokens:
        classes.append("is-centered")
    if "right" in tokens:
        classes.append("is-right")
    if "indented" in tokens:
        classes.append("is-indented")
    if "boxed" in tokens:
        classes.append("is-boxed")
    if "bold" in tokens:
        classes.append("is-bold")
    if "small" in tokens:
        classes.append("is-small")

    indent_level = min(max(safe_int(layout.get("indent_level"), 0), 0), 3)
    classes.append(f"indent-{indent_level}")

    text_size = str(layout.get("text_size", "")).strip().lower()
    if text_size in {"title", "subtitle", "heading", "body", "small", "tiny"}:
        classes.append(f"size-{text_size}")

    group = str(layout.get("group", "")).strip().lower()
    if group in {"header", "diagram", "body", "footer"}:
        classes.append(f"group-{group}")

    return " ".join(classes)


def render_element(element):
    """Turn one translated JSON element into HTML."""
    element_type = element.get("element_type", "body_block")
    classes = element_class(element)
    raw_text = element.get("slovenian_text", "").strip()

    if element_type == "chain_diagram":
        return render_chain_diagram(element, classes)

    if element_type == "divider":
        return f'<div class="{classes}" aria-hidden="true"></div>'

    if not raw_text:
        return ""

    marker = marker_for_element(element, raw_text)
    marker_types = {"list_item", "step", "subitem", "bullet_item"}
    if marker or element_type in marker_types:
        text = format_html_text(remove_repeated_marker(raw_text, marker))
        marker_html = html.escape(marker)
        return (
            f'<div class="{classes} marker-row">'
            f'<span class="marker">{marker_html}</span>'
            f'<div class="marker-content">{text}</div>'
            "</div>"
        )

    text = format_html_text(raw_text)

    if element_type == "title":
        return f'<h1 class="{classes}">{text}</h1>'
    if element_type == "subtitle":
        return f'<p class="{classes}">{text}</p>'
    if element_type == "heading":
        return f'<h2 class="{classes}">{text}</h2>'
    if element_type == "callout":
        return f'<div class="{classes}">{text}</div>'
    if element_type == "label":
        return f'<div class="{classes}">{text}</div>'
    if element_type == "diagram_text":
        return f'<span class="{classes}">{text}</span>'
    if element_type == "footer":
        return f'<footer class="{classes}">{text}</footer>'
    if element_type == "list_item":
        return f'<p class="{classes}">{text}</p>'
    return f'<p class="{classes}">{text}</p>'


def page_density_class(elements):
    """Use slightly tighter spacing when a scanned page contains a lot of text."""
    visible_text = " ".join(
        element.get("slovenian_text", "") for element in elements if element.get("slovenian_text")
    )
    if len(elements) > 34 or len(visible_text) > 3200:
        return "is-dense"
    return ""


def build_handout_html(result):
    """Create the recreated Slovenian handout as a clean A4 HTML page."""
    elements = result.get("page", {}).get("elements", [])
    elements = sorted(elements, key=lambda item: item.get("order", 0))
    density_class = page_density_class(elements)

    body_html = "\n".join(render_element(element) for element in elements)

    return f"""
<!doctype html>
<html lang="sl">
<head>
  <meta charset="utf-8">
  <style>
    @page {{
      size: A4;
      margin: 0;
    }}

    * {{
      box-sizing: border-box;
    }}

    html,
    body {{
      margin: 0;
      min-height: 100%;
      background: #eef0ef;
      font-family: Arial, Helvetica, sans-serif;
      color: #2b2b2b;
    }}

    body {{
      padding: 18px 0;
    }}

    .page {{
      width: 210mm;
      min-height: 297mm;
      margin: 0 auto;
      padding: 14mm 17mm 16mm;
      background: #ffffff;
      border: 1px solid #c9cecc;
      box-shadow: 0 8px 24px rgba(35, 43, 39, 0.14);
    }}

    .handout-element {{
      margin: 0 0 5.8px;
      line-height: 1.24;
      font-size: 10pt;
      page-break-inside: avoid;
      overflow-wrap: anywhere;
      hyphens: auto;
    }}

    .type-title {{
      text-align: center;
      font-size: 16.5pt;
      line-height: 1.08;
      font-weight: 700;
      letter-spacing: 0;
      margin-top: 1mm;
      margin-bottom: 4.5mm;
    }}

    .type-subtitle {{
      text-align: center;
      font-size: 9.4pt;
      line-height: 1.2;
      color: #4e4e4e;
      padding-bottom: 2.6mm;
      border-bottom: 1.2px solid #303030;
      margin-bottom: 5mm;
    }}

    .type-heading {{
      font-size: 11.4pt;
      line-height: 1.2;
      font-weight: 700;
      margin-top: 6.5mm;
      margin-bottom: 2.4mm;
    }}

    .type-body_block,
    .type-list_item,
    .type-step,
    .type-subitem,
    .type-bullet_item {{
      font-size: 9.9pt;
    }}

    .marker-row {{
      display: grid;
      grid-template-columns: 9mm minmax(0, 1fr);
      column-gap: 2mm;
      align-items: start;
      margin-bottom: 4.6px;
    }}

    .marker {{
      font-weight: 700;
      text-align: left;
      white-space: nowrap;
    }}

    .marker-content {{
      min-width: 0;
    }}

    .type-step {{
      margin-top: 3.8mm;
    }}

    .type-step .marker-content {{
      font-weight: 400;
    }}

    .type-subitem {{
      grid-template-columns: 7mm minmax(0, 1fr);
      margin-left: 7mm;
      font-size: 9.55pt;
    }}

    .type-bullet_item {{
      grid-template-columns: 5mm minmax(0, 1fr);
      margin-left: 7mm;
      font-size: 9.55pt;
    }}

    .type-label {{
      display: inline-block;
      margin: 2.2mm 7px 2.8mm 0;
      font-size: 8.9pt;
      font-weight: 700;
      text-transform: uppercase;
    }}

    .type-diagram_text {{
      display: inline-block;
      border: 1px solid #6f7471;
      border-radius: 999px;
      padding: 3px 8px;
      margin: 2mm 4px 3.2mm 0;
      font-size: 8.9pt;
      line-height: 1.1;
      font-weight: 700;
      background: #f6f7f6;
    }}

    .type-chain_diagram {{
      width: 100%;
      margin: 4mm 0 6mm;
      padding: 0;
    }}

    .chain-diagram svg {{
      display: block;
      width: 100%;
      height: auto;
      max-height: 44mm;
    }}

    .chain-oval {{
      fill: none;
      stroke: #3f4341;
      stroke-width: 3;
    }}

    .chain-oval-problem {{
      fill: #777777;
      fill-opacity: 0.45;
    }}

    .chain-arrow {{
      stroke: #4f5451;
      stroke-width: 2.2;
      stroke-linecap: round;
    }}

    .chain-arrow-head {{
      fill: #4f5451;
    }}

    .chain-label {{
      font-family: Arial, Helvetica, sans-serif;
      font-size: 18px;
      font-weight: 700;
      fill: #2f3331;
    }}

    .type-callout,
    .is-boxed {{
      border: 1px solid #bfc7c3;
      background: #f7f8f7;
      border-radius: 3px;
      padding: 2.6mm 3mm;
      margin: 3mm 0 3.6mm;
    }}

    .type-divider {{
      height: 0;
      border-top: 1.1px solid #303030;
      margin: 3mm 0 4mm;
    }}

    .type-footer {{
      margin-top: 6mm;
      padding-top: 2.5mm;
      border-top: 0.8px solid #c8c8c8;
      font-size: 8.1pt;
      color: #666666;
      text-align: right;
    }}

    .hint-section_break {{
      margin-top: 6.5mm;
    }}

    .hint-top_rule {{
      border-top: 1.1px solid #303030;
      padding-top: 2.4mm;
    }}

    .hint-bottom_rule {{
      border-bottom: 1.1px solid #303030;
      padding-bottom: 2.4mm;
      margin-bottom: 4.5mm;
    }}

    .hint-uppercase {{
      text-transform: uppercase;
    }}

    .hint-compact {{
      margin-bottom: 3px;
      line-height: 1.18;
    }}

    .is-centered {{
      text-align: center;
    }}

    .is-right {{
      text-align: right;
    }}

    .is-indented,
    .indent-1 {{
      margin-left: 7mm;
    }}

    .indent-2 {{
      margin-left: 14mm;
    }}

    .indent-3 {{
      margin-left: 21mm;
    }}

    .is-bold {{
      font-weight: 700;
    }}

    .is-small,
    .size-small {{
      font-size: 8.9pt;
    }}

    .size-tiny {{
      font-size: 8.1pt;
    }}

    .page.is-dense {{
      padding-top: 13mm;
      padding-bottom: 13mm;
    }}

    .page.is-dense .handout-element {{
      margin-bottom: 3.8px;
      line-height: 1.18;
    }}

    .page.is-dense .type-body_block,
    .page.is-dense .type-list_item,
    .page.is-dense .type-step,
    .page.is-dense .type-subitem,
    .page.is-dense .type-bullet_item {{
      font-size: 9.35pt;
    }}

    .page.is-dense .type-heading {{
      margin-top: 4.8mm;
      margin-bottom: 1.8mm;
      font-size: 10.8pt;
    }}

    @media print {{
      html,
      body {{
        width: 210mm;
        min-height: 297mm;
        background: #ffffff;
        padding: 0;
      }}

      .page {{
        width: 210mm;
        min-height: 297mm;
        margin: 0;
        border: 0;
        box-shadow: none;
      }}
    }}
  </style>
</head>
<body>
  <main class="page {density_class}">
    {body_html}
  </main>
</body>
</html>
""".strip()


def register_pdf_fonts():
    """Register Windows Arial fonts so Slovenian characters print correctly."""
    font_paths = {
        "HandoutRegular": "C:/Windows/Fonts/arial.ttf",
        "HandoutBold": "C:/Windows/Fonts/arialbd.ttf",
        "HandoutItalic": "C:/Windows/Fonts/ariali.ttf",
    }

    registered = {}
    for font_name, path in font_paths.items():
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, path))
                registered[font_name] = font_name
            except Exception:
                pass

    return {
        "regular": registered.get("HandoutRegular", "Helvetica"),
        "bold": registered.get("HandoutBold", "Helvetica-Bold"),
        "italic": registered.get("HandoutItalic", "Helvetica-Oblique"),
    }


class ChainDiagramFlowable(Flowable):
    """Draw the chain-analysis diagram in exported PDF files."""

    def __init__(self, labels, fonts):
        super().__init__()
        self.labels = labels
        self.fonts = fonts
        self.width = 170 * mm
        self.height = 42 * mm

    def wrap(self, avail_width, avail_height):
        self.width = avail_width
        return self.width, self.height

    def draw_rotated_ellipse(self, canvas, cx, cy, rx, ry, angle, fill=0):
        canvas.saveState()
        canvas.translate(cx, cy)
        canvas.rotate(angle)
        canvas.ellipse(-rx, -ry, rx, ry, stroke=1, fill=fill)
        canvas.restoreState()

    def draw_arrow_head(self, canvas, points):
        path = canvas.beginPath()
        path.moveTo(points[0][0], points[0][1])
        for x, y in points[1:]:
            path.lineTo(x, y)
        path.close()
        canvas.drawPath(path, stroke=0, fill=1)

    def draw(self):
        canvas = self.canv
        canvas.saveState()
        canvas.scale(self.width / 720, self.height / 180)

        canvas.setStrokeColor(colors.HexColor("#4f5451"))
        canvas.setFillColor(colors.HexColor("#4f5451"))
        canvas.setLineWidth(2.2)
        canvas.line(126, 70, 172, 78)
        self.draw_arrow_head(canvas, [(172, 78), (160, 72), (160, 85)])
        canvas.line(620, 54, 548, 68)
        self.draw_arrow_head(canvas, [(548, 68), (560, 60), (562, 73)])
        canvas.line(352, 132, 372, 96)
        self.draw_arrow_head(canvas, [(372, 96), (372, 110), (360, 104)])
        canvas.line(386, 132, 410, 96)
        self.draw_arrow_head(canvas, [(410, 96), (408, 110), (397, 103)])
        canvas.line(508, 126, 505, 92)
        self.draw_arrow_head(canvas, [(505, 92), (511, 105), (499, 105)])

        canvas.setStrokeColor(colors.HexColor("#3f4341"))
        canvas.setLineWidth(3)
        canvas.setFillColor(colors.white)
        for cx, cy, rx, ry, angle in [
            (208, 78, 58, 25, 10),
            (250, 94, 49, 22, 8),
            (318, 85, 39, 19, 8),
            (378, 88, 52, 25, 7),
            (438, 84, 47, 22, -6),
        ]:
            self.draw_rotated_ellipse(canvas, cx, cy, rx, ry, angle)

        canvas.setFillColor(colors.HexColor("#b3b3b3"))
        self.draw_rotated_ellipse(canvas, 505, 75, 78, 27, 1, fill=1)
        canvas.setFillColor(colors.white)
        for cx, cy, rx, ry, angle in [
            (590, 83, 36, 19, 8),
            (630, 101, 39, 19, 8),
            (662, 121, 41, 18, 11),
        ]:
            self.draw_rotated_ellipse(canvas, cx, cy, rx, ry, angle)

        canvas.setFillColor(colors.HexColor("#2f3331"))
        canvas.setFont(self.fonts["bold"], 17)
        canvas.drawCentredString(72, 66, self.labels["vulnerability"])
        canvas.drawCentredString(214, 132, self.labels["prompting_event"])
        canvas.drawCentredString(368, 151, self.labels["links"])
        canvas.drawCentredString(505, 132, self.labels["consequences"])
        canvas.drawCentredString(632, 48, self.labels["problem_behavior"])

        canvas.restoreState()


def pdf_safe_text(text):
    """Escape text for ReportLab paragraphs and preserve simple line breaks."""
    escaped = html.escape(text.strip())
    return escaped.replace("\n", "<br/>")


def get_pdf_styles():
    """Define simple A4 handout styles that match the HTML preview closely."""
    fonts = register_pdf_fonts()
    return {
        "_fonts": fonts,
        "title": ParagraphStyle(
            "HandoutTitle",
            fontName=fonts["bold"],
            fontSize=17,
            leading=21,
            alignment=TA_CENTER,
            spaceAfter=7,
        ),
        "subtitle": ParagraphStyle(
            "HandoutSubtitle",
            fontName=fonts["regular"],
            fontSize=10,
            leading=13,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4f4f4f"),
            borderWidth=0.8,
            borderColor=colors.HexColor("#333333"),
            borderPadding=(0, 0, 6, 0),
            spaceAfter=11,
        ),
        "heading": ParagraphStyle(
            "HandoutHeading",
            fontName=fonts["bold"],
            fontSize=12,
            leading=15,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "HandoutBody",
            fontName=fonts["regular"],
            fontSize=10.4,
            leading=13.3,
            alignment=TA_LEFT,
            spaceAfter=5,
        ),
        "body_bold": ParagraphStyle(
            "HandoutBodyBold",
            fontName=fonts["bold"],
            fontSize=9.9,
            leading=12.4,
            alignment=TA_LEFT,
            spaceAfter=5,
        ),
        "step": ParagraphStyle(
            "HandoutStep",
            fontName=fonts["regular"],
            fontSize=9.9,
            leading=12.4,
            alignment=TA_LEFT,
            spaceBefore=5,
            spaceAfter=4,
        ),
        "indented": ParagraphStyle(
            "HandoutIndented",
            fontName=fonts["regular"],
            fontSize=9.55,
            leading=12.2,
            alignment=TA_LEFT,
            leftIndent=7 * mm,
            spaceAfter=5,
        ),
        "callout": ParagraphStyle(
            "HandoutCallout",
            fontName=fonts["regular"],
            fontSize=10.2,
            leading=13,
            alignment=TA_LEFT,
            borderWidth=0.8,
            borderColor=colors.HexColor("#c8cfcc"),
            borderPadding=6,
            backColor=colors.HexColor("#f6f8f7"),
            spaceBefore=5,
            spaceAfter=8,
        ),
        "label": ParagraphStyle(
            "HandoutLabel",
            fontName=fonts["bold"],
            fontSize=9.5,
            leading=12,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#333333"),
            spaceAfter=4,
        ),
        "diagram_text": ParagraphStyle(
            "HandoutDiagramText",
            fontName=fonts["bold"],
            fontSize=9.5,
            leading=12,
            alignment=TA_LEFT,
            borderWidth=0.6,
            borderColor=colors.HexColor("#777777"),
            borderPadding=4,
            backColor=colors.HexColor("#f7f7f7"),
            spaceAfter=5,
        ),
        "footer": ParagraphStyle(
            "HandoutFooter",
            fontName=fonts["regular"],
            fontSize=8.5,
            leading=11,
            alignment=TA_RIGHT,
            textColor=colors.HexColor("#666666"),
            borderWidth=0.5,
            borderColor=colors.HexColor("#d0d0d0"),
            borderPadding=(6, 0, 0, 0),
            spaceBefore=12,
        ),
    }


def style_for_pdf_element(element, styles):
    """Map the model's element type and layout hint to a PDF paragraph style."""
    element_type = element.get("element_type", "body_block")
    tokens = layout_tokens_for_element(element)

    if element_type == "title":
        return styles["title"]
    if element_type == "subtitle":
        return styles["subtitle"]
    if element_type == "heading":
        return styles["heading"]
    if element_type == "callout" or "boxed" in tokens:
        return styles["callout"]
    if element_type == "label":
        return styles["label"]
    if element_type == "diagram_text":
        return styles["diagram_text"]
    if element_type == "footer":
        return styles["footer"]
    if element_type == "step":
        return styles["step"]
    if element_type in {"list_item", "subitem", "bullet_item"} or "indented" in tokens:
        return styles["indented"]
    if "bold" in tokens:
        return styles["body_bold"]
    return styles["body"]


def build_pdf_bytes(result):
    """Create a printable A4 PDF from the same structured elements as the preview."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=17 * mm,
        leftMargin=17 * mm,
        topMargin=14 * mm,
        bottomMargin=16 * mm,
        title="Slovenian handout",
    )

    styles = get_pdf_styles()
    elements = result.get("page", {}).get("elements", [])
    elements = sorted(elements, key=lambda item: item.get("order", 0))

    story = []
    for element in elements:
        if element.get("element_type") == "chain_diagram":
            story.append(ChainDiagramFlowable(diagram_labels_from_element(element), styles["_fonts"]))
            story.append(Spacer(1, 3 * mm))
            continue

        if element.get("element_type") == "divider":
            story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#303030")))
            story.append(Spacer(1, 2.5 * mm))
            continue

        text = export_text_for_element(element)
        if not text:
            continue

        story.append(Paragraph(pdf_safe_text(text), style_for_pdf_element(element, styles)))

        if element.get("element_type") in {"title", "heading", "callout"}:
            story.append(Spacer(1, 2.5 * mm))

    doc.build(story)
    return buffer.getvalue()


def count_pdf_pages(pdf_bytes):
    """Count PDF pages so the app can warn about multi-page output."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return len(reader.pages)


def set_docx_paragraph_bottom_border(paragraph):
    """Add a simple bottom border to a Word paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "333333")
    p_bdr.append(bottom)


def load_image_font(size, bold=False):
    """Load Arial for the Word diagram PNG, with a safe fallback."""
    font_path = "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"
    try:
        return ImageFont.truetype(font_path, size)
    except Exception:
        return ImageFont.load_default()


def draw_centered_image_text(draw, center_x, center_y, text, font, fill):
    """Draw centered text on a PIL image."""
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((center_x - width / 2, center_y - height / 2), text, font=font, fill=fill)


def draw_image_arrow(draw, start, end, fill):
    """Draw a simple arrow for the Word diagram image."""
    draw.line([start, end], fill=fill, width=4)
    end_x, end_y = end
    start_x, start_y = start
    direction_x = end_x - start_x
    direction_y = end_y - start_y
    length = max((direction_x**2 + direction_y**2) ** 0.5, 1)
    unit_x = direction_x / length
    unit_y = direction_y / length
    perp_x = -unit_y
    perp_y = unit_x
    back_x = end_x - unit_x * 18
    back_y = end_y - unit_y * 18
    points = [
        (end_x, end_y),
        (back_x + perp_x * 8, back_y + perp_y * 8),
        (back_x - perp_x * 8, back_y - perp_y * 8),
    ]
    draw.polygon(points, fill=fill)


def build_chain_diagram_png(labels):
    """Create a small PNG chain diagram for the editable Word export."""
    image = Image.new("RGB", (1100, 280), "white")
    draw = ImageDraw.Draw(image)
    label_font = load_image_font(26, bold=True)
    stroke = "#3f4341"
    arrow = "#4f5451"
    text = "#2f3331"

    draw_image_arrow(draw, (190, 105), (265, 116), arrow)
    draw_image_arrow(draw, (945, 82), (820, 105), arrow)
    draw_image_arrow(draw, (545, 207), (575, 150), arrow)
    draw_image_arrow(draw, (595, 207), (635, 150), arrow)
    draw_image_arrow(draw, (770, 195), (765, 142), arrow)

    for x1, y1, x2, y2, fill in [
        (270, 78, 420, 150, None),
        (345, 110, 490, 174, None),
        (485, 92, 600, 150, None),
        (585, 98, 740, 165, None),
        (690, 90, 830, 152, None),
        (735, 72, 960, 150, "#b3b3b3"),
        (890, 97, 995, 154, None),
        (940, 130, 1045, 188, None),
        (985, 158, 1090, 214, None),
    ]:
        draw.ellipse((x1, y1, x2, y2), outline=stroke, width=5, fill=fill)

    draw_centered_image_text(draw, 110, 96, labels["vulnerability"], label_font, text)
    draw_centered_image_text(draw, 355, 218, labels["prompting_event"], label_font, text)
    draw_centered_image_text(draw, 570, 245, labels["links"], label_font, text)
    draw_centered_image_text(draw, 765, 218, labels["consequences"], label_font, text)
    draw_centered_image_text(draw, 955, 62, labels["problem_behavior"], label_font, text)

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def add_docx_chain_diagram(doc, element):
    """Insert the visual chain diagram into the Word document."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run()
    run.add_picture(build_chain_diagram_png(diagram_labels_from_element(element)), width=Cm(16.2))


def add_docx_text(paragraph, text, bold=False, italic=False, color=None):
    """Add text to a Word paragraph with basic styling."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def configure_docx_paragraph(paragraph, element):
    """Apply alignment, indentation, and spacing based on element metadata."""
    element_type = element.get("element_type", "body_block")
    tokens = layout_tokens_for_element(element)
    layout = get_element_layout(element)

    paragraph.paragraph_format.space_after = Pt(3.8 if "compact" in tokens else 5)
    paragraph.paragraph_format.line_spacing = 1.08

    if "centered" in tokens or element_type in {"title", "subtitle"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif "right" in tokens or element_type == "footer":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    indent_level = safe_int(layout.get("indent_level"), 0)
    if element_type in {"list_item", "subitem", "bullet_item"}:
        indent_level = max(indent_level, 1)
    if indent_level > 0 or "indented" in tokens:
        paragraph.paragraph_format.left_indent = Cm(0.7 * max(indent_level, 1))

    if element_type in {"heading", "callout", "step"} or "section-break" in tokens:
        paragraph.paragraph_format.space_before = Pt(8)


def add_docx_element(doc, element):
    """Add one structured handout element to the editable Word document."""
    if element.get("element_type") == "chain_diagram":
        add_docx_chain_diagram(doc, element)
        return

    if element.get("element_type") == "divider":
        paragraph = doc.add_paragraph()
        set_docx_paragraph_bottom_border(paragraph)
        paragraph.paragraph_format.space_after = Pt(8)
        return

    text = export_text_for_element(element)
    if not text:
        return

    element_type = element.get("element_type", "body_block")
    tokens = layout_tokens_for_element(element)
    paragraph = doc.add_paragraph()
    configure_docx_paragraph(paragraph, element)

    bold = element_type in {"title", "heading", "label", "diagram_text"} or "bold" in tokens
    italic = "small" in tokens and element_type == "footer"

    run = add_docx_text(paragraph, text, bold=bold, italic=italic)

    if element_type == "title":
        run.font.size = Pt(17)
        paragraph.paragraph_format.space_after = Pt(8)
    elif element_type == "subtitle":
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(79, 79, 79)
        set_docx_paragraph_bottom_border(paragraph)
        paragraph.paragraph_format.space_after = Pt(12)
    elif element_type == "heading":
        run.font.size = Pt(12)
    elif element_type in {"label", "diagram_text"}:
        run.font.size = Pt(9.5)
    elif element_type in {"subitem", "bullet_item"}:
        run.font.size = Pt(9.5)
    elif element_type == "footer":
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(102, 102, 102)
        set_docx_paragraph_bottom_border(paragraph)
    else:
        run.font.size = Pt(9.9 if element_type == "step" else 10.0)

    if "bottom-rule" in tokens:
        set_docx_paragraph_bottom_border(paragraph)


def build_docx_bytes(result):
    """Create an editable A4 Word document from the reviewed elements."""
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for element in get_ordered_elements(result):
        add_docx_element(document, element)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_source_language_docx_bytes(result):
    """Create an editable Word document using the extracted source-language text."""
    source_result = source_language_result(result)
    return build_docx_bytes(source_result)


def source_language_result(result):
    """Return a copy of the result where export text is the source English text."""
    source_result = copy.deepcopy(result)
    for element in get_ordered_elements(source_result):
        english_text = element.get("english_text", "")
        if english_text:
            element["slovenian_text"] = english_text

        diagram = element.get("diagram")
        if isinstance(diagram, dict):
            labels = diagram.get("labels")
            if isinstance(labels, dict):
                for key, value in labels.items():
                    if isinstance(value, dict):
                        labels[key] = value.get("english_text") or value.get("text") or value.get("slovenian_text", "")
            elif isinstance(labels, list):
                for label in labels:
                    if isinstance(label, dict) and label.get("english_text"):
                        label["slovenian_text"] = label["english_text"]

    update_summary_fields(source_result)
    return source_result


def source_text_from_result(result):
    """Create plain source text from ordered English elements."""
    lines = []
    for element in get_ordered_elements(result):
        english_text = str(element.get("english_text") or "").strip()
        if english_text:
            lines.append(english_text)
    return "\n".join(lines)


def use_source_docx_as_main_input(result, docx_bytes):
    """Send the generated source-language DOCX to the main translation app."""
    st.session_state.source_text = source_text_from_result(result)
    st.session_state.source_file_name = make_source_download_name(result)
    st.session_state.source_file_type = "docx"
    st.session_state.source_file_bytes = docx_bytes

    for key, value in {
        "analysis_report": "",
        "translation_prompt": "",
        "translated_text": "",
        "proofread_text": "",
        "proofreading_baseline_text": "",
        "qa_report": "",
        "aligned_xliff_bytes": b"",
        "aligned_xliff_summary": "",
        "aligned_rows": [],
        "realigned_xliff_bytes": b"",
        "realigned_xliff_summary": "",
        "realigned_rows": [],
        "bilingual_review_rows": [],
        "reflow_summary": "",
        "net_word_grid_rows": [],
        "net_word_grid_summary": {},
        "cost_entries": [],
    }.items():
        st.session_state[key] = value

    st.session_state.handout_handoff_message = (
        f"Source DOCX from Handout Translator is now loaded under 1 Source: "
        f"{st.session_state.source_file_name}"
    )


def make_source_download_name(result):
    """Create a friendly filename for the source-language handout DOCX."""
    title = result.get("page", {}).get("title", {}).get("english_text", "")
    if not title:
        title = "source_handout"
    filename = re.sub(r"[^A-Za-z0-9_-]+", "_", title).strip("_")
    return f"{filename or 'source_handout'}_source.docx"


def make_download_name(result, extension):
    """Create a friendly download filename from the translated title."""
    title = result.get("page", {}).get("title", {}).get("slovenian_text", "")
    if not title:
        title = "slovenski_izrocek"

    filename = re.sub(r"[^A-Za-z0-9_-]+", "_", title).strip("_")
    return f"{filename or 'slovenski_izrocek'}.{extension}"


def get_ordered_elements(result):
    """Return page elements in their visual/order sequence."""
    elements = result.get("page", {}).get("elements", [])
    return sorted(elements, key=lambda item: item.get("order", 0))


def layout_hint_list(layout_hint):
    """Turn the model's layout hint string into checkbox-friendly tokens."""
    return split_layout_hint(layout_hint)


def layout_from_review_selection(element, selected_hints):
    """Keep the richer layout object aligned with the simple editor controls."""
    layout = copy.deepcopy(get_element_layout(element))
    selected = set(selected_hints)

    if "centered" in selected:
        layout["alignment"] = "center"
    elif "right" in selected:
        layout["alignment"] = "right"
    elif "left" in selected:
        layout["alignment"] = "left"
    else:
        layout.pop("alignment", None)

    if "bold" in selected:
        layout["emphasis"] = "bold"
    elif "uppercase" in selected:
        layout["emphasis"] = "uppercase"
    else:
        layout.pop("emphasis", None)

    if "small" in selected:
        layout["text_size"] = "small"
    elif "large" in selected:
        layout["text_size"] = "heading"
    else:
        layout.pop("text_size", None)

    if "numbered" in selected:
        layout["marker_style"] = "number"
    elif "lettered" in selected:
        layout["marker_style"] = "letter"
    elif "bullet" in selected:
        layout["marker_style"] = "bullet"
    else:
        layout.pop("marker_style", None)

    if "boxed" in selected:
        layout["border"] = "boxed"
    elif "top-rule" in selected:
        layout["border"] = "top-rule"
    elif "bottom-rule" in selected:
        layout["border"] = "bottom-rule"
    else:
        layout.pop("border", None)

    if "section-break" in selected:
        layout["spacing"] = "section-break"
    elif "compact" in selected:
        layout["spacing"] = "compact"
    else:
        layout.pop("spacing", None)

    if "indented" in selected:
        layout["indent_level"] = max(safe_int(layout.get("indent_level"), 0), 1)
    elif safe_int(layout.get("indent_level"), 0) > 0:
        layout.pop("indent_level", None)

    return layout


def update_summary_fields(result):
    """Keep title/subtitle/footer summary fields aligned with edited elements."""
    page = result.setdefault("page", {})

    for field in ["title", "subtitle", "footer"]:
        page.setdefault(field, {"english_text": "", "slovenian_text": ""})

    for element in get_ordered_elements(result):
        element_type = element.get("element_type")
        if element_type in {"title", "subtitle", "footer"}:
            page[element_type] = {
                "english_text": element.get("english_text", ""),
                "slovenian_text": element.get("slovenian_text", ""),
            }


def build_structured_from_elements(elements):
    """Rebuild the structured buckets from the edited ordered elements."""
    structured = {
        "title": [],
        "subtitle": [],
        "headings": [],
        "body_blocks": [],
        "callouts": [],
        "labels": [],
        "diagram_text": [],
        "diagrams": [],
        "footer": [],
    }

    bucket_by_type = {
        "title": "title",
        "subtitle": "subtitle",
        "heading": "headings",
        "body_block": "body_blocks",
        "list_item": "body_blocks",
        "step": "body_blocks",
        "subitem": "body_blocks",
        "bullet_item": "body_blocks",
        "callout": "callouts",
        "label": "labels",
        "diagram_text": "diagram_text",
        "chain_diagram": "diagrams",
        "footer": "footer",
    }

    for element in elements:
        bucket = bucket_by_type.get(element.get("element_type"))
        if bucket:
            if element.get("element_type") == "chain_diagram":
                structured[bucket].append(" -> ".join(diagram_labels_from_element(element).values()))
            else:
                structured[bucket].append(element.get("slovenian_text", ""))

    return structured


def build_sections_from_elements(elements):
    """Create review sections from edited elements so section view reflects edits."""
    sections = []
    current_title = "Prevedeni izroček"
    current_lines = []

    for element in elements:
        element_type = element.get("element_type")
        text = export_text_for_element(element)
        if not text:
            continue

        if element_type in {"title", "heading"}:
            if current_lines:
                sections.append({"title": current_title, "text": "\n".join(current_lines)})
                current_lines = []
            current_title = text
        else:
            current_lines.append(text)

    if current_lines:
        sections.append({"title": current_title, "text": "\n".join(current_lines)})

    return sections


def build_reviewed_result(result):
    """Collect the current review/edit form values into a new result object."""
    reviewed = copy.deepcopy(result)
    reviewed.setdefault("page", {})
    edited_elements = []
    result_version = st.session_state.get("result_version", 0)

    for index, element in enumerate(get_ordered_elements(result)):
        prefix = f"edit_{result_version}_{index}"
        include = st.session_state.get(f"{prefix}_include", True)
        if not include:
            continue

        edited_element = copy.deepcopy(element)
        edited_element["order"] = len(edited_elements) + 1
        edited_element["element_type"] = st.session_state.get(
            f"{prefix}_type",
            element.get("element_type", "body_block"),
        )
        edited_element["slovenian_text"] = st.session_state.get(
            f"{prefix}_slovenian",
            element.get("slovenian_text", ""),
        )
        selected_layout_hints = st.session_state.get(
            f"{prefix}_layout",
            layout_tokens_for_element(element),
        )
        edited_element["layout_hint"] = " ".join(selected_layout_hints)
        edited_element["layout"] = layout_from_review_selection(element, selected_layout_hints)
        edited_elements.append(edited_element)

    reviewed["page"]["elements"] = edited_elements
    reviewed["structured"] = build_structured_from_elements(edited_elements)
    reviewed["slovenian_sections"] = build_sections_from_elements(edited_elements)
    update_summary_fields(reviewed)
    return reviewed


def show_review_editor(result):
    """Let the user correct translated text and element types before export."""
    st.caption("Edit the Slovenian text and choose what each line is before exporting.")

    result_version = st.session_state.get("result_version", 0)
    elements = get_ordered_elements(result)

    if not elements:
        st.warning("No structured elements were found in the model result.")
        return

    for index, element in enumerate(elements):
        prefix = f"edit_{result_version}_{index}"
        element_type = element.get("element_type", "body_block")
        if element_type not in ELEMENT_TYPES:
            element_type = "body_block"

        english_text = element.get("english_text", "").strip()
        slovenian_text = element.get("slovenian_text", "").strip()

        label = f"{index + 1}. {element_type}"
        if slovenian_text:
            label = f"{label}: {slovenian_text[:70]}"

        with st.expander(label, expanded=index < 3):
            st.checkbox("Include in output", value=True, key=f"{prefix}_include")

            if english_text:
                st.text_area(
                    "Original English",
                    value=english_text,
                    height=85,
                    disabled=True,
                    key=f"{prefix}_english",
                )

            col_type, col_layout = st.columns([1, 2])
            with col_type:
                st.selectbox(
                    "Element type",
                    ELEMENT_TYPES,
                    index=ELEMENT_TYPES.index(element_type),
                    key=f"{prefix}_type",
                )
            with col_layout:
                st.multiselect(
                    "Layout hints",
                    LAYOUT_HINTS,
                    default=layout_tokens_for_element(element),
                    key=f"{prefix}_layout",
                )

            st.text_area(
                "Slovenian text",
                value=slovenian_text,
                height=120,
                key=f"{prefix}_slovenian",
            )


def show_translation_sections(result):
    """Display the Slovenian translation as reviewable sections."""
    for index, section in enumerate(result.get("slovenian_sections", []), start=1):
        title = section.get("title") or f"Section {index}"
        text = section.get("text") or ""

        with st.container(border=True):
            st.markdown(f"### {title}")
            st.write(text)


def show_result(result):
    """Display editor, layout preview, structured JSON, and translation sections."""
    result = normalize_result(copy.deepcopy(result))
    reviewed_result = build_reviewed_result(result)
    handout_html = build_handout_html(reviewed_result)
    pdf_bytes = build_pdf_bytes(reviewed_result)
    docx_bytes = build_docx_bytes(reviewed_result)
    source_docx_bytes = build_source_language_docx_bytes(reviewed_result)
    page_count = count_pdf_pages(pdf_bytes)

    tab_review, tab_preview, tab_json, tab_sections = st.tabs(
        ["Review/edit", "Recreated handout", "Structured JSON", "Translation sections"]
    )

    with tab_review:
        show_review_editor(result)
        st.info("Changes here update the preview and downloads after Streamlit refreshes the page.")
        with st.expander("Re-edit output with GPT-5.5", expanded=False):
            revision_instruction = st.text_area(
                "What should GPT-5.5 improve?",
                value=st.session_state.get(
                    "handout_revision_instruction",
                    "Make the layout more compact and ensure nothing is cut off in the A4 PDF.",
                ),
                height=110,
            )
            if st.button("Re-edit handout with GPT-5.5", type="primary", use_container_width=True):
                revise_handout_with_gpt(reviewed_result, revision_instruction, page_count)

    with tab_preview:
        st.caption("HTML/CSS preview with A4 PDF export.")

        if page_count == 1:
            st.success("PDF check: the handout fits on one A4 page.")
        else:
            st.warning(
                f"PDF check: this handout currently uses {page_count} A4 pages. "
                "That is okay for long handouts, but check the layout before printing."
            )

        html_bytes = handout_html.encode("utf-8")
        col_html, col_pdf, col_docx, col_source = st.columns(4)
        with col_html:
            st.download_button(
                "Download HTML",
                data=html_bytes,
                file_name=make_download_name(reviewed_result, "html"),
                mime="text/html",
            )
        with col_pdf:
            st.download_button(
                "Download A4 PDF",
                data=pdf_bytes,
                file_name=make_download_name(reviewed_result, "pdf"),
                mime="application/pdf",
            )
        with col_docx:
            st.download_button(
                "Download Word (.docx)",
                data=docx_bytes,
                file_name=make_download_name(reviewed_result, "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_source:
            st.download_button(
                "Download source Word",
                data=source_docx_bytes,
                file_name=make_source_download_name(reviewed_result),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            if st.button("Use source Word in 1 Source", use_container_width=True):
                use_source_docx_as_main_input(reviewed_result, source_docx_bytes)
                st.rerun()

        components.html(handout_html, height=950, scrolling=True)

    with tab_json:
        st.json(reviewed_result)

    with tab_sections:
        show_translation_sections(reviewed_result)

        unclear_text = reviewed_result.get("unclear_text", [])
        if unclear_text:
            with st.expander("Unclear or hard-to-read text"):
                for item in unclear_text:
                    st.write(f"- {item}")


def render_handout_translator_section():
    """Render the translator UI inside a larger Streamlit app."""
    st.title("Handout Translator")
    st.caption(
        "Current step: richer layout JSON → cleaner A4 handout recreation → PDF and Word export"
    )

    st.info(
        "This version extracts more layout detail, then recreates a cleaner A4 handout you can download as HTML, PDF, or Word."
    )

    api_key = get_api_key()
    if not api_key:
        st.error(
            "OpenAI API key is missing. Put it in `.streamlit/secrets.toml` "
            "before running the app."
        )
        st.stop()

    uploaded_image = st.file_uploader(
        "Upload a photo or scan of one English handout page",
        type=["png", "jpg", "jpeg"],
    )

    if uploaded_image is not None:
        st.image(uploaded_image, caption="Uploaded image", use_container_width=True)

    generate_button = st.button(
        "Extract, translate, and recreate layout",
        type="primary",
        disabled=uploaded_image is None,
    )

    if generate_button and uploaded_image is not None:
        image_data_url = uploaded_image_to_data_url(uploaded_image)

        with st.status("Reading image, translating, and recreating layout...", expanded=True) as status:
            try:
                st.write("Sending the image to OpenAI vision.")
                model_text = call_openai_vision(image_data_url)
                st.write("Parsing structured JSON.")
                result = parse_model_json(model_text)
            except RateLimitError:
                status.update(label="OpenAI quota problem.", state="error")
                st.error(
                    "OpenAI says your API quota or billing limit is exhausted. "
                    "The app code is working, but the API key cannot make this request right now."
                )
                return
            except json.JSONDecodeError:
                status.update(label="Finished, but JSON parsing failed.", state="error")
                st.warning("The model returned text that was not valid JSON. Showing raw output.")
                st.text_area("Raw model output", model_text, height=400)
                return
            except OpenAIError as error:
                status.update(label="OpenAI API error.", state="error")
                st.error(str(error))
                return

            status.update(label="Done.", state="complete")

        st.session_state["latest_result"] = result
        st.session_state["result_version"] = st.session_state.get("result_version", 0) + 1

    if "latest_result" in st.session_state:
        show_result(st.session_state["latest_result"])


def main():
    st.set_page_config(
        page_title="Handout Translator",
        page_icon="📄",
        layout="wide",
    )
    render_handout_translator_section()


if __name__ == "__main__":
    main()
