"""DOCX export helpers."""

from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from export_xliff import sentence_segments


WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
NAMESPACES = {"w": WORD_NAMESPACE}


def create_docx(title: str, body: str) -> bytes:
    """Create a DOCX file in memory and return it as bytes."""
    if not body.strip():
        raise ValueError("There is no text to export.")

    document = Document()
    document.add_heading(title, level=1)
    for paragraph in body.split("\n"):
        document.add_paragraph(paragraph)

    file_buffer = BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)
    return file_buffer.getvalue()


def create_bilingual_docx(source_text: str, target_text: str) -> bytes:
    """Create a source/target bilingual DOCX table segmented by sentence."""
    if not source_text.strip():
        raise ValueError("Source text is missing.")
    if not target_text.strip():
        raise ValueError("Target text is missing.")

    source_segments = sentence_segments(source_text)
    target_segments = sentence_segments(target_text)

    document = Document()
    document.add_heading("Bilingual Review Document", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Segment"
    headers[1].text = "Source"
    headers[2].text = "Target"

    for index, source in enumerate(source_segments, start=1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = source
        row[2].text = target_segments[index - 1] if index <= len(target_segments) else ""

    if len(target_segments) > len(source_segments):
        for extra_index, target in enumerate(target_segments[len(source_segments) :], start=len(source_segments) + 1):
            row = table.add_row().cells
            row[0].text = str(extra_index)
            row[1].text = ""
            row[2].text = target

    file_buffer = BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)
    return file_buffer.getvalue()


def create_bilingual_docx_from_rows(rows: list[dict[str, str]]) -> bytes:
    """Create a bilingual DOCX table from aligned rows."""
    if not rows:
        raise ValueError("No aligned rows are available.")

    document = Document()
    document.add_heading("Bilingual Review Document", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Segment"
    headers[1].text = "Source"
    headers[2].text = "Target"
    headers[3].text = "Confidence"
    headers[4].text = "Note"

    for index, row_data in enumerate(rows, start=1):
        row = table.add_row().cells
        row[0].text = str(row_data.get("id") or index)
        row[1].text = str(row_data.get("source") or "")
        row[2].text = str(row_data.get("target") or "")
        row[3].text = str(row_data.get("confidence") or "")
        row[4].text = str(row_data.get("note") or "")

    file_buffer = BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)
    return file_buffer.getvalue()


def read_bilingual_docx_review(docx_bytes: bytes) -> list[dict[str, str]]:
    """Read corrected target text from a bilingual review DOCX table.

    Reviewers are expected to edit the Target column. The Segment column is
    used to put corrections back into the matching in-app review rows.
    """
    if not docx_bytes:
        raise ValueError("Upload a corrected bilingual DOCX first.")

    document = Document(BytesIO(docx_bytes))
    for table in document.tables:
        if not table.rows:
            continue
        headers = [_cell_text(cell).lower() for cell in table.rows[0].cells]
        if "segment" not in headers or "target" not in headers:
            continue
        segment_col = headers.index("segment")
        target_col = headers.index("target")
        source_col = headers.index("source") if "source" in headers else None
        note_col = headers.index("note") if "note" in headers else None
        rows = []
        for table_row in table.rows[1:]:
            cells = table_row.cells
            segment = _cell_text(cells[segment_col]) if segment_col < len(cells) else ""
            target = _cell_text(cells[target_col]) if target_col < len(cells) else ""
            source = _cell_text(cells[source_col]) if source_col is not None and source_col < len(cells) else ""
            note = _cell_text(cells[note_col]) if note_col is not None and note_col < len(cells) else ""
            if segment or source or target:
                rows.append(
                    {
                        "Segment": segment,
                        "Source": source,
                        "Target": target,
                        "Review note": note,
                    }
                )
        if rows:
            return rows

    raise ValueError("No bilingual review table was found. Use the DOCX exported by this app.")


def create_formatted_docx_from_template(template_bytes: bytes, translated_text: str) -> bytes:
    """Replace text in the original DOCX while keeping the document XML.

    This keeps the original DOCX package, styles, fonts, images, tables,
    sections, margins, and run properties. Only text nodes in word/document.xml
    are changed. The translated text must keep one target paragraph/line for
    each extracted source paragraph/line.
    """
    if not template_bytes:
        raise ValueError("No original DOCX template is available.")
    if not translated_text.strip():
        raise ValueError("There is no translated text to export.")

    output = BytesIO()

    with ZipFile(BytesIO(template_bytes), "r") as source_zip:
        document_xml = source_zip.read("word/document.xml")
        root = DocumentXml.from_bytes(document_xml)
        paragraph_count = root.text_paragraph_count()
        target_segments = fit_text_to_paragraph_count(translated_text, paragraph_count)

        root.replace_paragraph_text(target_segments)
        updated_document_xml = root.to_bytes()

        with ZipFile(output, "w", ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                data = updated_document_xml if item.filename == "word/document.xml" else source_zip.read(item.filename)
                target_zip.writestr(item, data)

    output.seek(0)
    return output.getvalue()


def fit_text_to_paragraph_count(text: str, paragraph_count: int) -> list[str]:
    """Fit target text to the original DOCX paragraph count.

    This keeps DOCX formatting independent from bilingual sentence alignment.
    If the target already has the right number of non-empty lines, those lines
    are used. Otherwise sentence-like pieces are merged evenly into the source
    paragraph count.
    """
    if paragraph_count <= 0:
        raise ValueError("The source DOCX has no editable text paragraphs.")

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) == paragraph_count:
        return lines

    pieces = sentence_segments(text)
    if not pieces:
        raise ValueError("There is no translated text to export.")
    if len(pieces) == paragraph_count:
        return pieces
    if paragraph_count == 1:
        return [" ".join(pieces)]
    if len(pieces) < paragraph_count:
        return pieces + [""] * (paragraph_count - len(pieces))
    return _merge_evenly(pieces, paragraph_count)


class DocumentXml:
    """Tiny wrapper around Word document XML text replacement."""

    def __init__(self, root):
        self.root = root

    @classmethod
    def from_bytes(cls, document_xml: bytes) -> "DocumentXml":
        from xml.etree import ElementTree as ET

        ET.register_namespace("w", WORD_NAMESPACE)
        root = ET.fromstring(document_xml)
        return cls(root)

    def text_paragraph_count(self) -> int:
        """Count non-empty Word paragraphs in document order."""
        return len(self._text_paragraphs())

    def replace_paragraph_text(self, replacement_segments: list[str]) -> None:
        """Replace each text paragraph with the matching translated segment."""
        for paragraph, replacement in zip(self._text_paragraphs(), replacement_segments):
            text_nodes = paragraph.findall(".//w:t", NAMESPACES)
            _replace_text_nodes(text_nodes, replacement)

    def to_bytes(self) -> bytes:
        """Return updated XML bytes."""
        from xml.etree import ElementTree as ET

        return ET.tostring(self.root, encoding="utf-8", xml_declaration=True)

    def _text_paragraphs(self) -> list:
        paragraphs = []
        for paragraph in self.root.findall(".//w:p", NAMESPACES):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", NAMESPACES))
            if text.strip():
                paragraphs.append(paragraph)
        return paragraphs


def _replace_text_nodes(text_nodes: list, replacement: str) -> None:
    """Put translated text back into existing runs to preserve run formatting."""
    if not text_nodes:
        return

    original_lengths = [len(node.text or "") for node in text_nodes]
    total_length = sum(original_lengths)
    if total_length == 0:
        text_nodes[0].text = replacement
        return

    cursor = 0
    replacement_length = len(replacement)
    for index, node in enumerate(text_nodes):
        if index == len(text_nodes) - 1:
            chunk = replacement[cursor:]
        else:
            share = round((original_lengths[index] / total_length) * replacement_length)
            chunk = replacement[cursor : cursor + share]
            cursor += share
        node.text = chunk
        _set_space_preserve(node, chunk)


def _cell_text(cell) -> str:
    """Read all paragraph text from a Word table cell."""
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _set_space_preserve(node, text: str) -> None:
    """Tell Word to preserve leading/trailing spaces when a run needs them."""
    if text.startswith(" ") or text.endswith(" "):
        node.set(XML_SPACE, "preserve")
    elif XML_SPACE in node.attrib:
        del node.attrib[XML_SPACE]


def _merge_evenly(pieces: list[str], target_count: int) -> list[str]:
    """Merge many translated pieces into the original number of paragraphs."""
    merged = []
    total = len(pieces)
    for index in range(target_count):
        start = round(index * total / target_count)
        end = round((index + 1) * total / target_count)
        if end <= start:
            end = min(total, start + 1)
        merged.append(" ".join(pieces[start:end]).strip())
    return merged
