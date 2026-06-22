from __future__ import annotations

import unittest
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from export_docx import (
    create_bilingual_docx_from_rows,
    create_formatted_docx_from_template,
    fit_text_to_paragraph_count,
    read_bilingual_docx_review,
)


class ExportDocxTests(unittest.TestCase):
    def test_fit_text_to_paragraph_count_merges_alignment_segments(self):
        text = "One. Two. Three. Four."

        paragraphs = fit_text_to_paragraph_count(text, 2)

        self.assertEqual(len(paragraphs), 2)
        self.assertIn("One.", paragraphs[0])
        self.assertIn("Four.", paragraphs[1])

    def test_formatted_docx_export_keeps_source_paragraph_count(self):
        document = Document()
        document.add_paragraph("First source paragraph.")
        document.add_paragraph("Second source paragraph.")
        buffer = BytesIO()
        document.save(buffer)

        exported = create_formatted_docx_from_template(
            buffer.getvalue(),
            "First translated sentence. Extra sentence. Second translated sentence. Final sentence.",
        )
        exported_document = Document(BytesIO(exported))
        paragraphs = [paragraph.text for paragraph in exported_document.paragraphs if paragraph.text.strip()]

        self.assertEqual(len(paragraphs), 2)
        self.assertIn("First translated", paragraphs[0])
        self.assertIn("Final sentence", paragraphs[1])

    def test_bilingual_review_docx_round_trip_reads_corrected_target(self):
        docx_bytes = create_bilingual_docx_from_rows(
            [
                {
                    "id": "1",
                    "source": "Hello",
                    "target": "Pozdrav",
                    "confidence": 100,
                    "note": "Manual review row.",
                }
            ]
        )
        document = Document(BytesIO(docx_bytes))
        document.tables[0].rows[1].cells[2].text = "Popravljen pozdrav"
        corrected = BytesIO()
        document.save(corrected)

        rows = read_bilingual_docx_review(corrected.getvalue())

        self.assertEqual("1", rows[0]["Segment"])
        self.assertEqual("Popravljen pozdrav", rows[0]["Target"])

    def test_bilingual_review_docx_reads_tracked_insertions(self):
        docx_bytes = create_bilingual_docx_from_rows(
            [
                {
                    "id": "1",
                    "source": "Hello",
                    "target": "Old target",
                    "confidence": 100,
                    "note": "",
                }
            ]
        )
        corrected = _replace_docx_xml(
            docx_bytes,
            b"<w:t>Old target</w:t>",
            (
                b"<w:del><w:r><w:delText>Old target</w:delText></w:r></w:del>"
                b"<w:ins><w:r><w:t>New visible target</w:t></w:r></w:ins>"
            ),
        )

        rows = read_bilingual_docx_review(corrected)

        self.assertEqual("New visible target", rows[0]["Target"])


def _replace_docx_xml(docx_bytes: bytes, old: bytes, new: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(docx_bytes), "r") as source_zip:
        with ZipFile(output, "w", ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                data = source_zip.read(item.filename)
                if item.filename == "word/document.xml":
                    self_check = old in data
                    if not self_check:
                        raise AssertionError("Expected XML snippet was not found in test DOCX.")
                    data = data.replace(old, new, 1)
                target_zip.writestr(item, data)
    return output.getvalue()


if __name__ == "__main__":
    unittest.main()
